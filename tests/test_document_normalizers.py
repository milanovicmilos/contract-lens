"""Tests for PDF and DOCX normalizers — uses tempfile fixtures, no network."""

from pathlib import Path
from unittest.mock import patch

import docx
import pytest
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from src.data.document_normalizer import (
    DocumentFormat,
    DocumentNormalizerFactory,
    DocxDocumentNormalizer,
    PdfDocumentNormalizer,
)


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(path), pagesize=LETTER)
    c.drawString(72, 720, "DISTRIBUTION AGREEMENT")
    c.drawString(72, 700, "This agreement is governed by Delaware law.")
    c.showPage()
    c.drawString(72, 720, "Page 2 content: Non-Compete clause applies for 5 years.")
    c.showPage()
    c.save()
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = docx.Document()
    document.core_properties.title = "Sample Contract"
    document.core_properties.author = "Test"
    document.add_paragraph("This Agreement is governed by the laws of Delaware.")
    document.add_paragraph("The Distributor shall not compete worldwide.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Liability"
    table.rows[1].cells[1].text = "$10,000"
    document.save(path)
    return path


def test_pdf_normalizer_extracts_text_and_pages(pdf_file: Path):
    n = PdfDocumentNormalizer()
    doc = n.normalize(pdf_file)

    assert doc.format == DocumentFormat.PDF
    assert doc.page_count == 2
    assert "DISTRIBUTION AGREEMENT" in doc.content
    assert "Delaware" in doc.content
    assert "Non-Compete" in doc.content
    assert doc.metadata["page_count_actual"] == 2


def test_pdf_normalizer_raises_on_missing_file(tmp_path: Path):
    n = PdfDocumentNormalizer()
    with pytest.raises(FileNotFoundError):
        n.normalize(tmp_path / "nope.pdf")


def test_pdf_normalizer_raises_on_missing_pypdf(pdf_file: Path):
    with patch.dict("sys.modules", {"pypdf": None}):
        n = PdfDocumentNormalizer()
        with pytest.raises(RuntimeError, match="pypdf"):
            n.normalize(pdf_file)


def test_docx_normalizer_extracts_paragraphs_and_tables(docx_file: Path):
    n = DocxDocumentNormalizer()
    doc = n.normalize(docx_file)

    assert doc.format == DocumentFormat.DOCX
    assert "governed by the laws of Delaware" in doc.content
    assert "compete worldwide" in doc.content
    # Table cells should land as pipe-delimited rows.
    assert "Liability" in doc.content
    assert "$10,000" in doc.content
    assert doc.metadata.get("title") == "Sample Contract"
    assert doc.metadata.get("author") == "Test"


def test_docx_normalizer_raises_on_missing_file(tmp_path: Path):
    n = DocxDocumentNormalizer()
    with pytest.raises(FileNotFoundError):
        n.normalize(tmp_path / "nope.docx")


def test_factory_routes_by_suffix(pdf_file: Path, docx_file: Path):
    pdf_doc = DocumentNormalizerFactory.normalize(pdf_file)
    docx_doc = DocumentNormalizerFactory.normalize(docx_file)
    assert pdf_doc.format == DocumentFormat.PDF
    assert docx_doc.format == DocumentFormat.DOCX
